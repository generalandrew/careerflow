#!/usr/bin/env python3
"""
write_cover_letter.py - Generate a cover letter DOCX for an application.

Reads JSON on stdin:
  {
    "app_dir": "/abs/path",
    "company": "Acme",
    "role": "Senior Solutions Engineer",
    "candidate_name": "Andrew Richard Miller",
    "candidate_location": "Fort Worth, TX",
    "candidate_phone": "555-555-5555",
    "candidate_email": "you@example.com",
    "candidate_linkedin": "linkedin.com/in/...",
    "hiring_contact_name": "MacKenzie",                    (optional)
    "hiring_contact_title": "Talent Acquisition Manager",  (optional)
    "salutation": "Dear MacKenzie,",                       (optional, defaults to "Dear Hiring Manager,")
    "opening_paragraph": "I'm reaching out because...",
    "body_paragraphs": ["...", "...", "..."],
    "closing_paragraph": "I'd welcome the opportunity to...",
    "signoff": "Best,"                                     (optional, default "Best,")
  }

Writes <app_dir>/CoverLetter_<Company>.docx
"""
import sys, os, json, datetime, subprocess
from pathlib import Path


def slug(s):
    import re
    s = re.sub(r"[^A-Za-z0-9]+", "", s)
    return s[:40]


def main():
    p = json.load(sys.stdin)
    app_dir = Path(p["app_dir"]).resolve()
    company = p.get("company", "Company")
    out_docx = app_dir / f"CoverLetter_{slug(company)}.docx"

    # Build via docx module directly (Node not required for this one)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        sys.stderr.write("python-docx required: pip install python-docx\n"); return 2

    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    # Default font Calibri 11
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    today = datetime.date.today().strftime("%B %d, %Y")

    # Header block, contact info
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(p.get("candidate_name", ""))
    r.bold = True
    r.font.size = Pt(16)

    contact_bits = [
        p.get("candidate_location", ""),
        f"Phone: {p['candidate_phone']}" if p.get("candidate_phone") else "",
        f"Email: {p['candidate_email']}" if p.get("candidate_email") else "",
        p.get("candidate_linkedin", ""),
    ]
    contact_line = "  |  ".join([b for b in contact_bits if b])
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(contact_line)
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()  # spacer

    # Date
    dp = doc.add_paragraph(today)
    dp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # spacer

    # Hiring contact
    if p.get("hiring_contact_name"):
        hc = doc.add_paragraph()
        hc.add_run(p["hiring_contact_name"])
        if p.get("hiring_contact_title"):
            hc.add_run(", " + p["hiring_contact_title"])
        hc.add_run(f"\n{company}")
    else:
        doc.add_paragraph(f"Hiring Team\n{company}")

    doc.add_paragraph()  # spacer

    # Salutation
    salutation = p.get("salutation", "Dear Hiring Manager,")
    doc.add_paragraph(salutation)

    doc.add_paragraph()  # spacer

    # Body
    if p.get("opening_paragraph"):
        doc.add_paragraph(p["opening_paragraph"])
    for b in p.get("body_paragraphs", []):
        doc.add_paragraph(b)
    if p.get("closing_paragraph"):
        doc.add_paragraph(p["closing_paragraph"])

    doc.add_paragraph()  # spacer

    # Sign-off
    doc.add_paragraph(p.get("signoff", "Best,"))
    doc.add_paragraph(p.get("candidate_name", ""))

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    print(f"Wrote {out_docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
